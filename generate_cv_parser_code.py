import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# --- CONFIGURATION ---

# The specific folder for the CV Parser code
TARGET_FOLDER = 'cv-parser'

# Name of the output Word document
OUTPUT_FILENAME = 'Appendix_D_CVParser_Code.docx'

# Directories to strictly IGNORE
EXCLUDE_DIRS = {
    'venv', '.venv', 'env',          # Virtual Environments
    '__pycache__',                   # Python Cache
    '.git', '.idea', '.vscode',      # IDE & Git
    'node_modules', 'dist', 'build'  # Misc
}

# File extensions/names to INCLUDE
INCLUDE_EXTS = {
    '.py',                           # Python Source
    'requirements.txt',              # Dependencies
    'Dockerfile',                    # Container Config
    '.env.example'                   # Config examples
}

def create_cv_parser_doc():
    doc = Document()
    
    # Document Title
    heading = doc.add_heading('Appendix D: CV Parser Implementation', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add a brief intro note
    intro = doc.add_paragraph('This appendix contains the source code for the AI-powered CV Parsing microservice built with Python and FastAPI.')
    intro.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('_' * 80)

    base_path = os.path.join('.', TARGET_FOLDER)
    
    if not os.path.exists(base_path):
        print(f"Error: Directory '{TARGET_FOLDER}' not found in the current folder.")
        print("Make sure you place this script next to the 'cv-parser' folder.")
        return

    files_processed = 0

    # Walk through the directory
    for root, dirs, files in os.walk(base_path):
        # Modify dirs in-place to skip excluded directories
        dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
        
        for file in files:
            # Check extension or exact filename (like Dockerfile)
            file_ext = os.path.splitext(file)[1]
            
            if (file_ext in INCLUDE_EXTS or file in INCLUDE_EXTS):
                file_path = os.path.join(root, file)
                
                print(f"Processing: {file_path}")
                
                try:
                    # Add file path header (Styled Dark Blue)
                    p = doc.add_heading(file_path, level=2)
                    p.style.font.color.rgb = RGBColor(0, 51, 102)
                    p.style.font.size = Pt(11)

                    # Read content
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Add Code Block with Courier Font
                    code_para = doc.add_paragraph(content)
                    code_para.style.font.name = 'Courier New'
                    code_para.style.font.size = Pt(9)
                    code_para.paragraph_format.space_after = Pt(2)
                    
                    # Add Separator
                    doc.add_paragraph('_' * 80)
                    files_processed += 1

                except Exception as e:
                    print(f"Error reading {file_path}: {e}")

    if files_processed > 0:
        doc.save(OUTPUT_FILENAME)
        print(f"\n✅ Success! Generated {OUTPUT_FILENAME} with {files_processed} files.")
    else:
        print(f"\n⚠️ No matching files found in '{TARGET_FOLDER}'. Check your folder structure.")

if __name__ == '__main__':
    create_cv_parser_doc()