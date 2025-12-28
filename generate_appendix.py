import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# --- CONFIGURATION ---

# Define the folders you want to turn into separate documents
# Format: ('Folder Name', 'Output Filename.docx')
TARGET_FOLDERS = [
    ('backend', 'Appendix_D_Backend_Code.docx'),
    ('frontend', 'Appendix_D_Frontend_Code.docx'),
    ('cv-parser', 'Appendix_D_CVParser_Code.docx'),
]

# Directories to IGNORE globally
EXCLUDE_DIRS = {
    'node_modules', '.git', '.idea', '.vscode', 'dist', 'build', 
    'coverage', '__pycache__', 'migrations', 'public', 'assets', 
    'uploads', 'docs', '.claude'  # <--- Added 'docs' here
}

# File extensions to INCLUDE
INCLUDE_EXTS = {
    '.ts', '.tsx', '.js', '.jsx',  # JS/TS
    '.py',                         # Python
    '.prisma', '.sql',             # DB
    '.yml', '.yaml',               # Configs
    '.conf', '.sh', '.dockerignore', 'Dockerfile' # Infrastructure
}

# Files to specifically IGNORE
EXCLUDE_FILES = {
    'package-lock.json', 'yarn.lock', 'pnpm-lock.yaml', 
    'tsconfig.tsbuildinfo', 'vite.svg', 'hrflow.svg', '.DS_Store'
}

def create_doc_from_folder(source_folder, output_filename, root_only=False):
    """
    Scans a specific folder and writes its code content to a Word doc.
    if root_only is True, it only scans files in that directory, not subdirectories.
    """
    doc = Document()
    
    # Document Title
    heading = doc.add_heading(f'Implementation Details: {source_folder if source_folder != "." else "Root Infrastructure"}', 0)
    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    base_path = os.path.join('.', source_folder)
    
    if not os.path.exists(base_path):
        print(f"Skipping {source_folder}: Directory not found.")
        return

    files_processed = 0

    # Walk through the directory
    for root, dirs, files in os.walk(base_path):
        # Modify dirs in-place to skip excluded directories
        dirs[:] = [d for d in dirs if d not in EXCLUDE_DIRS]
        
        # If we only want root files (for docker-compose etc), skip subdirectories
        if root_only and root != base_path:
            continue

        for file in files:
            # Check extension or exact filename (like Dockerfile)
            file_ext = os.path.splitext(file)[1]
            
            if (file_ext in INCLUDE_EXTS or file in INCLUDE_EXTS) and file not in EXCLUDE_FILES:
                file_path = os.path.join(root, file)
                
                print(f"[{source_folder}] Adding: {file_path}")
                
                try:
                    # Add file path header (Styled Blue)
                    p = doc.add_heading(file_path, level=2)
                    p.style.font.color.rgb = RGBColor(0, 51, 102)
                    p.style.font.size = Pt(10)

                    # Read content
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()

                    # Add Code Block
                    code_para = doc.add_paragraph(content)
                    code_para.style.font.name = 'Courier New'
                    code_para.style.font.size = Pt(8)
                    code_para.paragraph_format.space_after = Pt(2)
                    
                    # Add Separator
                    doc.add_paragraph('_' * 100)
                    files_processed += 1

                except Exception as e:
                    print(f"Error reading {file_path}: {e}")

    if files_processed > 0:
        doc.save(output_filename)
        print(f"✅ Successfully created {output_filename} ({files_processed} files)")
    else:
        print(f"⚠️ No matching files found in {source_folder}")

if __name__ == '__main__':
    # 1. Generate docs for specific folders
    for folder, filename in TARGET_FOLDERS:
        create_doc_from_folder(folder, filename)

    # 2. Generate a doc for root-level files (Docker compose, etc.)
    # We pass root_only=True so it doesn't scan into backend/frontend again
    create_doc_from_folder('.', 'Appendix_D_Infrastructure.docx', root_only=True)